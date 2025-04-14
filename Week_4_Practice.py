import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
# -----------------------------
# Scenario 1: Merging students and grades (inner join: only students with recorded grades)
students = pd.DataFrame({
    'student_id': [1, 2, 3, 4],
    'name': ['Alice', 'Bob', 'Charlie', 'Diana']
})
grades = pd.DataFrame({
    'student_id': [1, 2, 4],
    'exam_score': [88, 92, 79]
})
# Inner join: keeps only students with recorded grades
students_with_grades = pd.merge(students, grades, on='student_id', how='inner')
print("Students with grades:")
print(students_with_grades)
print("\n")

# -----------------------------
# Scenario 2: Merging employee and department data 
# (ensuring all departments are shown even if no employees belong to them)
employees = pd.DataFrame({
    'emp_id': [101, 102],
    'name': ['John Doe', 'Jane Smith'],
    'dept_id': [1, 2]
})
departments = pd.DataFrame({
    'dept_id': [1, 2, 3],
    'dept_name': ['Sales', 'Engineering', 'HR']
})
# Use a right join (or left join, with departments as primary)
dept_full = pd.merge(employees, departments, on='dept_id', how='right')
print("All departments with their employees (if available):")
print(dept_full)
print("\n")

# -----------------------------
# Scenario 3: Combining Spring, Summer, and Fall admissions vertically
spring = pd.DataFrame({'student_id': [1, 2], 'Name': ['Spring','Akbar']})
summer = pd.DataFrame({'student_id': [3], 'Name': ['Summer']})
fall   = pd.DataFrame({'student_id': [4, 5], 'Name': ['Fall', 'Faris']})
all_admissions = pd.concat([spring, summer, fall], axis=0, ignore_index=True)
print("Combined Admissions:")
print(all_admissions)
print("\n")

# -----------------------------
# Scenario 4: Merging two DataFrames and losing rows from the first one
# (This happens with an inner join when some rows do not have matching keys)
df1 = pd.DataFrame({'key': [1,2,3,4], 'value1': ['A','B','C','D']})
df2 = pd.DataFrame({'key': [2,4], 'value2': ['W','X']})
merged_inner = pd.merge(df1, df2, on='key', how='inner')
print("Result of inner join (some rows lost):")
print(merged_inner)
print("\n")

# -----------------------------
# Scenario 5: Total sales per product category using groupby
df_sales = pd.DataFrame({
    'Category': ['Electronics', 'Clothing', 'Electronics', 'Furniture'],
    'Sales': [1000, 500, 700, 1200]
})
sales_summary = df_sales.groupby('Category')['Sales'].sum()
print("Sales Summary:")
print(sales_summary)
print("\n")

# -----------------------------
# Scenario 6: Average rating per customer using groupby
df_reviews = pd.DataFrame({
    'Customer': ['Alice', 'Bob', 'Alice', 'Bob'],
    'Rating': [4, 5, 3, 4]
})
avg_rating = df_reviews.groupby('Customer')['Rating'].mean()
print("Average Rating per Customer:")
print(avg_rating)
print("\n")

# -----------------------------
# Scenario 7: Count of students in grade brackets using value_counts
df_scores = pd.DataFrame({
    'Student': ['A', 'B', 'C', 'D', 'E'],
    'Grade': ['A', 'B', 'C', 'A', 'B']
})
grade_counts = df_scores['Grade'].value_counts()
print("Grade Counts:")
print(grade_counts)
print("\n")

# -----------------------------
# Scenario 8: Grouping continuous age values into brackets using pd.cut
df_age = pd.DataFrame({'Age': [5, 17, 25, 36, 55, 80]})
bins = [0, 18, 35, 60, 100]
labels = ['0-18', '19-35', '36-60', '60+']
df_age['Age_Group'] = pd.cut(df_age['Age'], bins=bins, labels=labels)
print("Age Groups:")
print(df_age)
print("\n")

# -----------------------------
# Scenario 9: Categorizing temperatures using apply with a custom function
df_temp = pd.DataFrame({'Celsius': [5, 20, 30, 35, 15]})
def categorize_temp(c):
    if c < 15:
        return 'Cold'
    elif c < 30:
        return 'Warm'
    else:
        return 'Hot'
df_temp['Category'] = df_temp['Celsius'].apply(categorize_temp)
print("Temperature Categories:")
print(df_temp)
print("\n")

# -----------------------------
# Scenario 10: Applying a 10% discount using vectorized arithmetic
df_prices = pd.DataFrame({'Price': [100, 250, 75]})
df_prices['Discounted'] = df_prices['Price'] * 0.9
print("Prices with Discount:")
print(df_prices)
print("\n")

# -----------------------------
# Scenario 11: Renaming columns to lowercase using list comprehension
df_columns = pd.DataFrame({'Name': ['John', 'Jane'], 'Age': [28, 34]})
df_columns.columns = [col.lower() for col in df_columns.columns]
print("Columns renamed to lowercase:")
print(df_columns.columns)
print("\n")

# -----------------------------
# Scenario 12: Saving a DataFrame to CSV without an index
# For demonstration, we will simply show the command (file not actually saved here)
# df_columns.to_csv('processed_data.csv', index=False)

# -----------------------------
# Scenario 13: Visualizing the distribution of a numeric column (Income) using a histogram
df_income = pd.DataFrame({'Income': [25000, 40000, 50000, 60000, 75000]})
plt.figure(figsize=(8, 5))
sns.histplot(df_income['Income'], bins=5)
plt.title('Income Distribution')
plt.xlabel('Income')
plt.ylabel('Frequency')
plt.show()
# Interpretation:
# The histogram shows how incomes are distributed across the sample.
# For instance, if the bars are relatively equal in height, income distribution is uniform; 
# if one bar is significantly higher, many individuals fall into that income range.

# -----------------------------
# Scenario 14: Comparing monthly sales of two products using a line plot
df_monthly_sales = pd.DataFrame({
    'Month': ['Jan', 'Feb', 'Mar', 'Apr'],
    'Product_A': [100, 150, 200, 170],
    'Product_B': [90, 140, 210, 180]
})
plt.figure(figsize=(8, 5))
df_monthly_sales.set_index('Month').plot(marker='o')
plt.title('Monthly Sales Comparison')
plt.ylabel('Sales')
plt.xlabel('Month')
plt.show()
# Interpretation:
# The line plot compares trends for two products. Crossovers or divergence in lines 
# may indicate shifts in customer preference over time.

# -----------------------------
# Scenario 15: Scatter plot for relationship between Math and Science scores
df_scores_subjects = pd.DataFrame({
    'Math': [80, 90, 70, 85],
    'Science': [75, 88, 68, 90]
})
plt.figure(figsize=(8, 5))
plt.scatter(df_scores_subjects['Math'], df_scores_subjects['Science'])
plt.title('Math vs Science Scores')
plt.xlabel('Math Scores')
plt.ylabel('Science Scores')
plt.show()
# Interpretation:
# A scatter plot allows us to visually assess the relationship between math and science scores.
# A positive trend would indicate that higher math scores tend to coincide with higher science scores.

# -----------------------------
# Scenario 16: Histogram and KDE curve for customer ages using Seaborn’s histplot
df_customers = pd.DataFrame({'Age': [20, 25, 30, 35, 40, 45, 50]})
plt.figure(figsize=(8, 5))
sns.histplot(df_customers['Age'], bins=7, kde=True)
plt.title('Customer Age Distribution with KDE')
plt.xlabel('Age')
plt.ylabel('Frequency')
plt.show()
# Interpretation:
# The combination of histogram and KDE shows both the raw counts and the smooth estimation of age distribution,
# giving a more complete picture of demographic spread.

# -----------------------------
# Scenario 17: Jointplot (scatter + histograms) for two numerical variables
# In this example, we will reuse Math and Science scores.
sns.jointplot(data=df_scores_subjects, x='Math', y='Science', kind='scatter', height=6)
plt.suptitle('Joint Distribution: Math vs Science', y=1.02)
plt.show()
# Interpretation:
# The jointplot provides a combined view: the central scatter plot shows the relationship,
# while the top and right histograms display the individual distributions of Math and Science scores.

# -----------------------------
# Scenario 18: Comparing median income across job categories using a boxplot
df_jobs = pd.DataFrame({
    'Job': ['Engineer', 'Doctor', 'Teacher', 'Engineer', 'Doctor', 'Teacher'],
    'Income': [70000, 80000, 50000, 72000, 82000, 53000]
})
plt.figure(figsize=(8, 5))
sns.boxplot(data=df_jobs, x='Job', y='Income')
plt.title('Income Distribution by Job')
plt.xlabel('Job Category')
plt.ylabel('Income')
plt.show()
# Interpretation:
# The boxplot shows the median, quartile ranges, and possible outliers of incomes within each job category.
# This allows us to compare the central tendency and spread among jobs.

# -----------------------------
# Scenario 19: Count plot for survey responses by gender
df_survey = pd.DataFrame({
    'Gender': ['Male', 'Female', 'Male', 'Female', 'Female'],
    'Product_Pref': ['A', 'B', 'A', 'C', 'B']
})
plt.figure(figsize=(8, 5))
sns.countplot(data=df_survey, x='Product_Pref', hue='Gender')
plt.title('Product Preferences by Gender')
plt.xlabel('Product Preference')
plt.ylabel('Count')
plt.show()
# Interpretation:
# The count plot breaks down responses (product preferences) by gender.
# Differences in the height of bars suggest varying popularity of products among genders.

# -----------------------------
# Scenario 20: Visualizing a correlation matrix with a heatmap using color intensity
df_features = pd.DataFrame({
    'Feature1': [1, 2, 3, 4, 5],
    'Feature2': [5, 4, 3, 2, 1],
    'Feature3': [2, 3, 2, 3, 2]
})
corr_matrix = df_features.corr()
plt.figure(figsize=(6, 5))
sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
plt.title('Correlation Heatmap')
plt.show()

# Interpretation:
# In the correlation heatmap, each cell’s color indicates the strength and direction 
# of the linear relationship between two features. Darker (or more intense) colors 
# closer to 1 or -1 denote strong positive or negative correlations respectively.
# For example, if Feature1 and Feature2 show a value of -1, that implies a perfect negative correlation.

# -----------------------------
# Generate a report file (Word document) that summarizes the graph interpretations.
doc = Document()
doc.add_heading('Pandas and Seaborn Analysis Report', 0)

doc.add_heading('Correlation Heatmap Interpretation', level=1)
doc.add_paragraph(
    "The correlation heatmap above visualizes the pairwise correlations among features. "
    "The annotated values indicate the Pearson correlation coefficients. "
    "Values close to 1 or -1 (with intense colors) show strong positive or negative linear relationships, "
    "while values near 0 (lighter colors) indicate weaker relationships. "
    "For instance, a perfect negative correlation (-1) between Feature1 and Feature2 suggests that as Feature1 increases, "
    "Feature2 decreases proportionally. Such visualizations are useful for identifying redundant features and understanding feature interactions."
)

doc.add_heading('Other Graph Interpretations', level=1)
doc.add_paragraph(
    "1. The Income Distribution histogram shows the spread of income data, revealing whether incomes are uniformly "
    "distributed or concentrated in a particular range.\n"
    "2. The Monthly Sales line plot indicates the trends over months for different products, aiding in detecting peaks, dips, or crossovers.\n"
    "3. The Math vs Science scatter plot allows identification of any correlation between the two exam scores.\n"
    "4. The Jointplot combines scatter and marginal histograms, providing both relationship insights and individual "
    "distribution details.\n"
    "5. The Boxplot by Job Category exposes the median, quartiles, and any outliers, facilitating comparison across job groups.\n"
    "6. The Count Plot for Product Preferences by Gender clearly shows the counts of responses split by gender."
)

# Save the Word document
report_path = "Pandas_Seaborn_Analysis_Report.docx"
doc.save(report_path)
print(f"\nReport generated and saved as '{report_path}'.")
